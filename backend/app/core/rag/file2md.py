"""
文档 → Markdown 转换器
支持 PDF / DOCX / XLSX / TXT / CSV / MD
PDF 和 DOCX 中的图片会提取并上传到 MinIO
"""
import io
import os
import uuid
import logging
from typing import Optional

import fitz  # PyMuPDF
import docx
from lxml import etree
import pandas as pd

from minio import Minio

logger = logging.getLogger(__name__)


class File2Markdown:
    """将任意文档转为 Markdown，图片存 MinIO"""

    def __init__(self, minio_client: Minio, bucket: str, images_prefix: str = ""):
        self.minio = minio_client
        self.bucket = bucket
        self.images_prefix = images_prefix  # e.g. "kb_name/images"

    # ======================== 统一入口 ========================

    def convert(self, file_bytes: bytes, file_type: str, filename: str = "") -> str:
        """
        将文件字节转为 Markdown 字符串
        返回: markdown 文本
        """
        converters = {
            "pdf": self._pdf_to_md,
            "docx": self._docx_to_md,
            "xlsx": self._xlsx_to_md,
            "txt": self._txt_to_md,
            "csv": self._csv_to_md,
            "md": self._passthrough_md,
        }
        fn = converters.get(file_type)
        if not fn:
            raise ValueError(f"Unsupported file type: {file_type}")
        return fn(file_bytes)

    # ======================== PDF ========================

    def _pdf_to_md(self, data: bytes) -> str:
        doc = fitz.open(stream=data, filetype="pdf")

        # 统计字体大小 → 识别正文 vs 标题
        font_counter: dict[float, int] = {}
        for page in doc:
            for block in page.get_text("dict")["blocks"]:
                if "lines" not in block:
                    continue
                for line in block["lines"]:
                    for span in line["spans"]:
                        sz = round(span["size"], 1)
                        font_counter[sz] = font_counter.get(sz, 0) + 1

        if not font_counter:
            return ""

        # 出现最多的字号 = 正文
        body_size = max(font_counter, key=font_counter.get)
        # 比正文大的字号 → 映射为 # 标题层级
        larger_sizes = sorted([s for s in font_counter if s > body_size], reverse=True)
        font2heading = {sz: "#" * (i + 1) for i, sz in enumerate(larger_sizes[:5])}

        parts: list[str] = []
        for page_num, page in enumerate(doc):
            blocks = page.get_text("dict")["blocks"]
            for block in blocks:
                if "image" in block:
                    # 提取图片 → 上传 MinIO
                    img_url = self._upload_image(block["image"], block.get("ext", "png"))
                    if img_url:
                        parts.append(f"\n![image]({img_url})\n")
                elif "lines" in block:
                    line_texts = []
                    for line in block["lines"]:
                        for span in line["spans"]:
                            text = span["text"].strip()
                            if not text:
                                continue
                            sz = round(span["size"], 1)
                            if sz in font2heading:
                                line_texts.append(f"{font2heading[sz]} {text}")
                            else:
                                line_texts.append(text)
                    if line_texts:
                        parts.append(" ".join(line_texts))
            parts.append("")  # 页面分隔

        doc.close()
        return "\n\n".join(parts)

    # ======================== DOCX ========================

    def _docx_to_md(self, data: bytes) -> str:
        doc = docx.Document(io.BytesIO(data))
        paragraphs = list(doc.paragraphs)
        tables = list(doc.tables)

        # 提取所有图片 rId → MinIO URL
        images: dict[str, str] = {}
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                img_bytes = rel.target_part.blob
                ext = os.path.splitext(rel.target_part.partname)[-1].lstrip(".")
                url = self._upload_image(img_bytes, ext)
                if url:
                    images[rel.rId] = url

        md_blocks: list[str] = []
        for block in doc.element.body:
            tag = block.tag if isinstance(block.tag, str) else ""
            if tag.endswith("p") and paragraphs:
                para = paragraphs.pop(0)
                md_blocks.append(self._docx_para_to_md(para, images))
            elif tag.endswith("tbl") and tables:
                table = tables.pop(0)
                md_blocks.append(self._docx_table_to_md(table))

        return "\n\n".join(md_blocks)

    def _docx_para_to_md(self, para, images: dict) -> str:
        style = para.style.name if para.style else ""
        prefix = ""
        if "List" in style:
            level = self._docx_list_level(para)
            prefix = "  " * level + "- "
        elif "Heading" in style:
            try:
                level = int(style.split(" ")[1])
            except (IndexError, ValueError):
                level = 1
            prefix = "#" * level + " "

        text = self._docx_parse_run(para, images)
        return prefix + text

    def _docx_parse_run(self, element, images: dict) -> str:
        parts = []
        for child in element.iter_inner_content():
            if isinstance(child, str):
                parts.append(child)
            elif isinstance(child, docx.text.run.Run):
                t = self._docx_parse_run(child, images)
                if child.bold:
                    t = f"**{t.strip()}**"
                if child.italic:
                    t = f"*{t.strip()}*"
                parts.append(t)
            elif isinstance(child, docx.text.hyperlink.Hyperlink):
                parts.append(f"[{child.text}]({child.address})")
            elif isinstance(child, docx.drawing.Drawing):
                rId = self._extract_r_embed(child._element.xml)
                if rId and rId in images:
                    parts.append(f"![image]({images[rId]})")
        return " ".join(parts)

    @staticmethod
    def _extract_r_embed(xml_str: str) -> Optional[str]:
        root = etree.fromstring(xml_str)
        ns = {
            "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
            "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        }
        blip = root.find(".//a:blip", namespaces=ns)
        if blip is not None:
            return blip.attrib.get(
                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
            )
        return None

    @staticmethod
    def _docx_list_level(para) -> int:
        p = para._element
        numPr = p.find(".//w:numPr", namespaces=p.nsmap)
        if numPr is not None:
            ilvl = numPr.find(".//w:ilvl", namespaces=p.nsmap)
            if ilvl is not None:
                return int(
                    ilvl.get(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
                    )
                )
        return 0

    @staticmethod
    def _docx_table_to_md(table) -> str:
        rows = []
        for i, row in enumerate(table.rows):
            cells = " | ".join(cell.text.strip() for cell in row.cells)
            rows.append(f"| {cells} |")
            if i == 0:
                rows.append("| " + " | ".join("---" for _ in row.cells) + " |")
        return "\n".join(rows)

    # ======================== XLSX ========================

    def _xlsx_to_md(self, data: bytes) -> str:
        xls = pd.ExcelFile(io.BytesIO(data))
        parts = []
        for sheet in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet)
            df = df.dropna(axis=1, how="all")
            if df.empty:
                continue
            try:
                table_md = df.to_markdown(index=False)
            except Exception:
                table_md = self._df_to_md(df)
            parts.append(f"### {sheet}\n\n{table_md}")
        return "\n\n".join(parts)

    @staticmethod
    def _df_to_md(df) -> str:
        header = "| " + " | ".join(str(c) for c in df.columns) + " |"
        sep = "| " + " | ".join("---" for _ in df.columns) + " |"
        rows = []
        for _, row in df.iterrows():
            rows.append("| " + " | ".join(str(v) for v in row) + " |")
        return "\n".join([header, sep] + rows)

    # ======================== TXT ========================

    def _txt_to_md(self, data: bytes) -> str:
        return data.decode("utf-8", errors="replace")

    # ======================== CSV ========================

    def _csv_to_md(self, data: bytes) -> str:
        df = pd.read_csv(io.StringIO(data.decode("utf-8", errors="replace")))
        df = df.dropna(axis=1, how="all")
        if df.empty:
            return ""
        try:
            return df.to_markdown(index=False)
        except Exception:
            return self._df_to_md(df)

    # ======================== MD (passthrough) ========================

    def _passthrough_md(self, data: bytes) -> str:
        return data.decode("utf-8", errors="replace")

    # ======================== 图片上传 ========================

    def _upload_image(self, image_data: bytes, ext: str = "png") -> Optional[str]:
        """上传图片到 MinIO，返回 presigned URL (不含签名，走后端代理)"""
        try:
            ext = ext.lower().lstrip(".")
            if ext not in ("png", "jpg", "jpeg", "gif", "bmp", "webp", "svg"):
                ext = "png"
            img_id = str(uuid.uuid4())
            object_name = f"{self.images_prefix}/{img_id}.{ext}" if self.images_prefix else f"images/{img_id}.{ext}"

            self.minio.put_object(
                self.bucket,
                object_name,
                io.BytesIO(image_data),
                length=len(image_data),
                content_type=f"image/{ext}",
            )
            # 返回相对路径，前端通过后端 API 代理访问
            return f"/api/v1/files/{self.bucket}/{object_name}"
        except Exception as e:
            logger.warning(f"Failed to upload image to MinIO: {e}")
            return None
